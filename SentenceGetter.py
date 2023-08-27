import nltk
from tika import parser
import os
import re
import docx
from textblob import TextBlob 
import pkg_resources
from symspellpy.symspellpy import SymSpell
keywords = ["problem", "challeng", "difficult", "protest", "demanding", "strenous", "limit", "hinder", "issue", "concern", "conflict", "imposs", "struggl", "tens", " not ", "need", "unclear", "never", "preven"]
sanitize = ['/',"@","\n"]
forbidden = [".com","www.","https:"]
word = True


def check_file():
    extension = "txt"
    if word:
        extension="docx"
    source = "D:\McGill OneDrive\OneDrive - McGill University\Side Projects\RA Work\Barnett\OPR\Challenge Searching"
    exists = os.path.exists(source + "\Results\output."+extension)
    if exists:
        counter = 0;
        while exists:
            filename = "output"+str(counter)+"."+extension
            fullpath = source +"\Results\\" + filename
            exists = os.path.exists(fullpath)
            counter +=1
    else:
        fullpath = source +"\Results\output."+extension
    if word:
        return fullpath, docx.Document()
    return fullpath, open(fullpath, "w+", encoding="utf-8")


def grab_text(output):
    first = True
    files= list()   
    for (path, currentDirectory, filenames) in os.walk("D:\McGill OneDrive\OneDrive - McGill University\Side Projects\RA Work\Barnett\OPR\Challenge Searching\Working"):
        for file in filenames:
            files += [os.path.join(path, file)]  
    for file in files:
        parsed_pdf = parser.from_file(file)
        data = parsed_pdf['content']
        for string in sanitize:
            data = data.replace(string, "")
        tokens = nltk.sent_tokenize(data)
        sentences = []
        for sentence in tokens:
            sentence = re.sub("\n", "", sentence)
            sentence = re.sub("-", "", sentence)
            sentence = (sym_spell.word_segmentation(sentence)).corrected_string
            sentence = str(TextBlob(sentence))
            if any(ele in sentence.lower() for ele in keywords) and not any(ele in sentence.lower() for ele in forbidden):
                sentences.append(sentence)
        name = file[file.rfind("\\")+1:]
        writer(output,word,first, name, sentences)
        first = False

def writer(output, word, first, name, sentences):
    if word:
        if first:
                output.add_paragraph(name.upper())
                first = not first
        else:
            output.add_paragraph("\n \n" + name.upper())
        for selected in sentences:
            output.add_paragraph("\n" + selected)
    else:
        if first:
            output.write(name.upper())
            first = not first
        else:
            output.write("\n \n" + name.upper())
        for selected in sentences:
            output.write("\n" + selected)


def main():
        fullpath, output = check_file()
        grab_text(output)
        if word:    
            output.save(fullpath)



if __name__=="__main__":
    sym_spell = SymSpell(max_dictionary_edit_distance=0, prefix_length=7)
    dictionary_path = pkg_resources.resource_filename("symspellpy", "frequency_dictionary_en_82_765.txt")
    sym_spell.load_dictionary(dictionary_path, term_index=0, count_index=1)
    main();